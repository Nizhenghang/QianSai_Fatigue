from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r'D:\Nizhenghang\QianSai\outputs\RA8D1_驾驶员疲劳检测系统_实验报告.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
BORDER = '9EADBC'


def set_run_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para(p, before=0, after=6, line=1.15, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    set_para(p, after=0, line=1.1)
    r = p.add_run(text)
    set_run_font(r, size=10, bold=bold)
    if fill:
        shade_cell(cell, fill)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_para(p, before=14 if level == 1 else 10, after=7)
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, name='黑体', size=16, bold=True, color=DARK_BLUE)
    else:
        set_run_font(r, name='黑体', size=13, bold=True, color=BLUE)
    return p


def add_para(doc, text, bold=False, align=None):
    p = doc.add_paragraph()
    set_para(p, after=6, line=1.18, align=align)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=bold)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        set_para(p, after=3, line=1.15)
        p.clear()
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = Inches(widths[i])
        set_cell_text(hdr[i], h, bold=True, fill=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Inches(widths[i])
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(val), align=align)
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is None:
                tcMar = OxmlElement('w:tcMar')
                tcPr.append(tcMar)
            for m, w in [('top','80'),('bottom','80'),('start','120'),('end','120')]:
                node = tcMar.find(qn(f'w:{m}'))
                if node is None:
                    node = OxmlElement(f'w:{m}')
                    tcMar.append(node)
                node.set(qn('w:w'), w)
                node.set(qn('w:type'), 'dxa')
    doc.add_paragraph()
    return table


def add_placeholder(doc, caption, hint, height_rows=4):
    table = doc.add_table(rows=height_rows, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    for r in range(1, height_rows):
        cell = cell.merge(table.cell(r, 0))
    cell.width = Inches(5.8)
    shade_cell(cell, 'F4F8FC')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=0, line=1.2)
    r = p.add_run('[图片预留]\n')
    set_run_font(r, name='微软雅黑', size=12, bold=True, color=BLUE)
    r = p.add_run(caption + '\n')
    set_run_font(r, name='微软雅黑', size=11, bold=True, color='404040')
    r = p.add_run(hint)
    set_run_font(r, name='微软雅黑', size=10, color='606060')
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(cap, before=2, after=8, line=1.0)
    rr = cap.add_run('图：' + caption)
    set_run_font(rr, size=10, color='555555')


def add_page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.5)
    sec.footer_distance = Inches(0.5)

    styles = doc.styles
    styles['Normal'].font.name = '宋体'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    styles['Normal'].font.size = Pt(11)

    header = sec.header.paragraphs[0]
    header.text = 'RA8D1 端侧视觉 AI 驾驶员疲劳检测系统实验报告'
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header.runs[0], size=9, color='666666')

    footer = sec.footer.paragraphs[0]
    footer.text = '实验报告'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color='666666')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(title, before=24, after=14, line=1.2)
    r = title.add_run('驾驶员疲劳检测系统实验报告')
    set_run_font(r, name='黑体', size=20, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=5)
    r = p.add_run('作品名称：RA8D1 端侧视觉 AI 驾驶员疲劳检测系统')
    set_run_font(r, name='宋体', size=12, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=18)
    r = p.add_run('作者信息：________________    日期：________________')
    set_run_font(r, name='宋体', size=10, color='666666')

    add_heading(doc, '摘要', 1)
    add_para(doc, '本作品面向驾驶安全场景，设计并实现了一套基于 RA8D1 单片机的驾驶员疲劳检测系统。系统以 MT9V03X 摄像头作为图像采集入口，在单片机端完成驾驶员面部区域的裁剪、缩放和轻量级 CNN 推理，并通过连续疲劳分数累积机制判断当前状态是否处于正常、轻度疲劳、警告疲劳、严重疲劳或无人脸状态。与仅依赖单次动作触发的方案相比，本作品将单帧图像识别结果和时间维度上的状态变化结合起来，降低了短暂眨眼、低头或偶发误判带来的误报风险。')
    add_para(doc, '系统硬件部分主要由 RA8D1 开发板、MT9V03X 摄像头、外接扬声器模块、HC-05 蓝牙模块及手机端显示 APP 组成。软件部分包含图像采集、ROI 预处理、CNN 前向推理、疲劳等级判断、语音播报、蓝牙数据发送和 Android 手机端显示等模块。模型训练流程由 PC 端采集脚本、数据增强脚本、训练脚本和模型导出脚本组成，可将训练得到的权重导出为 C 语言数组并部署到单片机固件中。')
    add_para(doc, '目前系统已经实现端侧实时检测、疲劳等级显示、严重疲劳语音报警、无人脸提醒以及手机端状态同步等功能。该作品具有端侧部署、低成本硬件、交互方式直观和便于后续扩展等特点，可用于智能驾驶辅助、疲劳驾驶预警、嵌入式 AI 教学展示及低成本安全监测实验平台。')

    add_heading(doc, '第一部分  作品概述', 1)
    add_heading(doc, '1.1 功能与特性', 2)
    add_para(doc, '本作品的核心功能是对驾驶员状态进行实时监测，并根据疲劳程度进行分级提示。系统启动后，摄像头持续采集驾驶员前方图像，固件从图像中裁剪出固定 ROI 区域并缩放为 64×32 灰度输入，随后调用轻量级 CNN 模型完成单帧正常/疲劳分类。后处理模块不会直接依据一次分类结果报警，而是将连续帧的疲劳判断转化为疲劳分数，并根据阈值映射为正常、轻度、警告和严重疲劳等级。')
    add_bullets(doc, ['端侧实时推理：不依赖云端传输图像，隐私性和独立性较好。', '疲劳等级分级：输出 NORMAL、LIGHT、WARN、DANGER 和 NO_FACE 状态。', '语音播报提醒：通过外接扬声器模块播放不同疲劳等级的提示语音。', '蓝牙手机显示：通过 HC-05 模块将状态数据发送到 Android APP。', '数据闭环：支持图像采集、数据增强、模型训练和权重导出。'])
    add_placeholder(doc, '作品整体实物图预留', '建议放置正面或斜 45°拍摄的整机照片。')

    add_heading(doc, '1.2 应用领域', 2)
    add_para(doc, '该系统主要应用于驾驶安全与嵌入式智能感知场景。在实际交通环境中，疲劳驾驶容易导致反应迟缓、注意力下降甚至交通事故，因此低成本、实时、可独立运行的疲劳检测装置具有较强的应用价值。本作品也适合作为嵌入式 AI 教学平台，用于展示从数据采集、模型训练到 MCU 部署的完整流程。')
    add_para(doc, '除车载疲劳检测外，该方案还可扩展到学习状态监测、值班人员精神状态提醒、工业岗位安全监测等场景。在这些场景中，系统可以结合不同摄像头安装角度和阈值策略，对长时间低头、闭眼、离岗等状态进行提示。')
    add_placeholder(doc, '典型应用场景图预留', '可放置车内安装位置、驾驶员测试场景或应用示意图。', 3)

    add_heading(doc, '1.3 主要技术特点', 2)
    add_para(doc, '本作品采用“端侧视觉 AI + 时序状态判断 + 多通道提醒”的设计思路。AI 模型不是部署在 PC 或云端，而是以 C 数组和手写前向推理的形式集成在 RA8D1 固件中。模型结构为 64×32×1 输入，经两层卷积、两层池化和两层全连接输出二分类结果。疲劳检测模块进一步结合置信度阈值、连续疲劳帧、疲劳分数累积和目标丢失检测完成最终状态判断。')
    add_bullets(doc, ['手写 CNN 前向推理，便于在无通用推理框架的 MCU 上运行。', '中间特征图放置于 SDRAM，减轻片上存储压力。', '采用固定 ROI 和灰度图输入，降低计算量。', '通过蓝牙串口协议输出状态帧，便于手机端显示和调试。', '语音提醒替代单一蜂鸣器，使报警信息更直观。'])

    add_heading(doc, '1.4 主要性能指标', 2)
    add_para(doc, '系统当前以功能实现和现场演示为主，部分量化指标需要结合后续统一测试数据进一步完善。已确认的固件构建信息和关键参数如下。')
    add_table(doc, ['指标项', '当前值或说明'], [
        ['主控平台', 'RA8D1 开发板'],
        ['图像输入', 'MT9V03X 摄像头，ROI 裁剪后缩放为 64×32 灰度图'],
        ['模型结构', 'Conv(5×5,8,stride2) + Pool + Conv(3×3,16) + Pool + FC(32) + FC(2)'],
        ['状态等级', 'NORMAL / LIGHT / WARN / DANGER / NO_FACE'],
        ['蓝牙通信', 'HC-05 经典蓝牙 SPP，UART3 9600 baud'],
        ['固件构建', 'Keil 构建 0 Error(s), 0 Warning(s)'],
        ['准确率/误报率', '预留：待统一测试集统计后填写'],
    ], [1.7, 4.6])

    add_heading(doc, '1.5 主要创新点', 2)
    add_bullets(doc, ['将轻量 CNN 模型部署到单片机端，实现不依赖 PC/云端的疲劳识别。', '采用连续疲劳分数而非单帧触发，能够更符合真实疲劳检测的时间连续性。', '将语音播报和手机 APP 显示结合，使报警结果更易理解和展示。', '形成自采数据、训练、导出、固件部署的完整嵌入式 AI 闭环。'])

    add_heading(doc, '1.6 设计流程', 2)
    add_para(doc, '设计流程从需求分析开始，首先确定系统需要完成驾驶员状态采集、疲劳识别、报警提醒和手机显示。随后完成硬件选型与连接，构建 RA8D1、摄像头、扬声器和蓝牙模块组成的实验平台。软件层面先实现图像采集和导出，再采集正常与疲劳样本，使用训练脚本得到 CNN 模型并导出为固件权重。最后在固件中实现推理、连续状态判断、语音播报和蓝牙数据发送，并通过实测不断调整疲劳阈值。')
    add_placeholder(doc, '系统设计流程图预留', '可放置从需求分析到模型部署的流程图或手绘图。', 3)

    add_page_break(doc)
    add_heading(doc, '第二部分  系统组成及功能说明', 1)
    add_para(doc, '本部分从整体结构、硬件系统和软件系统三个层次说明作品的设计细节。系统由图像采集层、端侧推理层、状态判断层和交互输出层组成，各层之间通过固定数据流连接，最终完成从驾驶员图像到疲劳等级提示的闭环。')

    add_heading(doc, '2.1 整体介绍', 2)
    add_para(doc, '系统整体框图可概括为：MT9V03X 摄像头采集图像，RA8D1 读取图像缓冲区并裁剪 ROI，CNN 推理模块输出正常/疲劳分类结果，疲劳检测模块进行连续分数累积和状态判断，最后由语音播报模块与 HC-05 蓝牙模块分别完成本地提醒和手机端显示。')
    add_placeholder(doc, '系统整体框图预留', '建议放置“摄像头→RA8D1→CNN→状态判断→语音/蓝牙”的框图。', 4)

    add_heading(doc, '2.2 硬件系统介绍', 2)
    add_para(doc, '硬件系统以 RA8D1 开发板为核心。MT9V03X 摄像头负责采集驾驶员图像，外接扬声器模块用于播放语音提示，HC-05 蓝牙模块通过 UART3 与主控通信，将疲劳等级和分数发送到手机端。系统供电和信号连接需要保证公共地可靠，蓝牙模块 RXD 接 RA8D1 UART3 TX，引脚电平保持 3.3 V TTL 兼容。')
    add_placeholder(doc, '硬件整体连接照片预留', '建议放置 RA8D1、摄像头、扬声器、HC-05 的整体连接照片。', 4)
    add_para(doc, '机械结构方面，当前作品以实验平台形式实现，摄像头需要固定在驾驶员正前方或略偏上的位置，以保证 ROI 区域能够覆盖眼部、口部和低头动作特征。后续可增加固定支架或外壳，提高比赛现场演示的稳定性。')
    add_placeholder(doc, '摄像头固定方式或外壳结构图预留', '可放置支架、安装位置、外壳草图或实物照片。', 3)
    add_para(doc, '电路模块方面，重点包括摄像头接口、扬声器输出模块和蓝牙串口通信模块。摄像头数据进入 RA8D1 的图像采集路径，语音模块通过本地音频播放逻辑输出不同等级提醒，蓝牙模块则发送 ASCII 状态帧，例如 BT|STATE|level=...|status=...|score=...。')
    add_placeholder(doc, '关键电路或接线图预留', '建议放置 HC-05 接线、扬声器模块接线或开发板局部原理图。', 3)

    add_heading(doc, '2.3 软件系统介绍', 2)
    add_para(doc, '软件系统包括固件端和手机端两部分。固件端运行在 RA8D1 上，主入口完成 SDRAM、摄像头、调试串口和相关外设初始化，然后在主循环中调用 fatigue_process_frame() 处理每一帧图像。手机端为 Android 蓝牙 APP，用于连接 HC-05 并解析来自单片机的状态帧。')
    add_placeholder(doc, '软件整体流程图预留', '建议放置固件主循环、模型推理、蓝牙发送、APP 显示的流程图。', 4)
    add_para(doc, '固件核心函数流程如下：首先判断摄像头帧是否完成；随后裁剪 FATIGUE_ROI_LEFT 至 FATIGUE_ROI_RIGHT、FATIGUE_ROI_TOP 至 FATIGUE_ROI_BOTTOM 区域，并缩放到 CNN_INPUT_W=64、CNN_INPUT_H=32；接着调用 cnn_classify() 得到疲劳分类和置信度；最后更新投票窗口、疲劳分数、报警状态和 HMI/蓝牙显示。')
    add_para(doc, '模型训练流程位于 training 目录。collect_frames.py 用于串口采集样本，augment_data.py 用于数据增强，train_cnn.py 用于训练 Keras CNN，export_model.py 将模型权重导出到 E18_usb_cdc_demo/code/cnn_weights.h。固件编译后即可在单片机端运行新的模型。')
    add_placeholder(doc, '训练与部署流程图预留', '可放置 collect_frames → train_cnn → export_model → Keil 固件的流程图。', 3)

    add_page_break(doc)
    add_heading(doc, '第三部分  完成情况及性能参数', 1)
    add_heading(doc, '3.1 整体介绍', 2)
    add_para(doc, '目前作品已经完成从硬件搭建、模型训练导出、固件推理、疲劳等级判断、语音报警到蓝牙手机显示的主要功能闭环。系统能够在检测到无人脸时给出提示，在疲劳分数达到严重等级时触发语音报警，并将当前状态同步发送至手机 APP。')
    add_placeholder(doc, '系统整体实物正面照片预留', '请放置整个系统正面照片。', 4)
    add_placeholder(doc, '系统整体实物斜 45°照片预留', '请放置整体斜 45°照片，用于展示装配效果。', 4)

    add_heading(doc, '3.2 工程成果', 2)
    add_para(doc, '硬件成果包括 RA8D1 主控平台、摄像头采集模块、扬声器语音提醒模块和 HC-05 蓝牙通信模块。软件成果包括嵌入式固件、训练脚本、模型权重文件和 Android 手机端 APP。')
    add_placeholder(doc, '硬件实物局部照片预留', '建议展示摄像头、蓝牙模块、扬声器模块等局部。', 3)
    add_placeholder(doc, '手机 APP 界面照片预留', '建议展示 NORMAL、LIGHT、WARN、DANGER 或 NO_FACE 状态界面。', 3)
    add_placeholder(doc, '语音报警或现场演示照片预留', '建议放置严重疲劳报警时的现场照片。', 3)

    add_heading(doc, '3.3 特性成果', 2)
    add_para(doc, '系统已经实现的主要特性包括端侧图像推理、连续疲劳等级判断、无人脸检测提醒、语音播报和蓝牙状态同步。当前固件最近一次构建结果为 0 Error(s)、0 Warning(s)，生成 HEX 文件可直接烧录到目标板进行演示。')
    add_table(doc, ['测试项目', '现象或结果', '备注'], [
        ['正常状态', '手机端显示 NORMAL，语音不报警', '需放置正常状态测试照片'],
        ['轻度疲劳', '疲劳分数上升后显示 LIGHT', '短暂动作不应频繁触发'],
        ['严重疲劳', '达到阈值后语音播报严重疲劳提醒', '需放置报警照片或视频截图'],
        ['无人脸状态', '长时间无有效目标时进入 NO_FACE', '可展示离开摄像头区域'],
        ['蓝牙通信', 'HC-05 向 APP 发送状态帧', 'APP 截图待插入'],
        ['模型准确率', '待统一测试集统计', '后续补充混淆矩阵'],
    ], [1.25, 3.0, 2.0])
    add_placeholder(doc, '性能测试或串口日志截图预留', '可放置构建日志、串口输出、APP 数据变化或测试记录。', 3)

    add_heading(doc, '第四部分  总结', 1)
    add_heading(doc, '4.1 可扩展之处', 2)
    add_para(doc, '后续可以从四个方向继续优化：第一，扩充自采数据集，覆盖不同光照、姿态、距离和驾驶员个体差异；第二，建立独立测试集，统计准确率、召回率、误报率和混淆矩阵；第三，优化手机端界面，将疲劳分数曲线、历史报警记录和连接状态可视化；第四，设计固定支架和外壳，提高系统在比赛现场演示时的稳定性和美观度。')

    add_heading(doc, '4.2 心得体会', 2)
    add_para(doc, '本作品的实现过程体现了嵌入式 AI 项目从算法到工程落地的完整链路。最初的难点并不只是训练一个能够区分正常与疲劳图像的模型，而是如何让模型在单片机上稳定运行，并将单帧识别结果转化为符合真实场景的连续疲劳判断。实际调试中发现，单次打哈欠、短暂低头和摄像头范围丢失都可能影响判断结果，因此系统逐步加入了置信度阈值、连续分数累积、无人脸检测和报警冷却等机制。')
    add_para(doc, '在硬件调试方面，串口屏方案曾受供电和通信细节影响，最终改用 HC-05 蓝牙与手机 APP 显示，使作品展示更加稳定。语音报警部分也从单纯蜂鸣器输出升级为扬声器语音提示，提升了交互效果。通过这些调整可以体会到，比赛作品不仅要“能识别”，还要“能稳定演示、能解释清楚、能让观众理解”。')
    add_para(doc, '通过本项目，我进一步理解了数据采集、模型训练、模型导出、嵌入式部署和现场验证之间的关系。后续如果继续完善，应重点补足数据统计和实验验证部分，让作品不仅具备功能完整性，也具备更充分的工程证据。')

    add_heading(doc, '第五部分  参考文献', 1)
    for ref in ['[1] Renesas Electronics. RA8D1 Group User Manual: Hardware.', '[2] Renesas Electronics. Flexible Software Package Documentation.', '[3] SEEKFREE. RA8D1 开源库与逐飞相关外设驱动资料.', '[4] TensorFlow/Keras Documentation. Convolutional Neural Network Model Training.', '[5] Android Developers. Bluetooth Classic and Serial Communication Documentation.', '[6] HC-05 Bluetooth Module User Manual.']:
        add_para(doc, ref)

    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    build()

