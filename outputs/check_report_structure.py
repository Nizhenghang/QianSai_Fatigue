from pathlib import Path
from docx import Document
p = Path(r'D:\Nizhenghang\QianSai\outputs\report_ascii_check.docx')
doc = Document(p)
texts = [x.text for x in doc.paragraphs if x.text.strip()]
cell_texts = [cell.text for t in doc.tables for row in t.rows for cell in row.cells]
all_text = '\n'.join(texts + cell_texts)
print('paragraphs', len(texts))
print('tables', len(doc.tables))
print('placeholders', all_text.count('[图片预留]'))
for key in ['摘要','第一部分  作品概述','第二部分  系统组成及功能说明','第三部分  完成情况及性能参数','第四部分  总结','第五部分  参考文献']:
    print(key, key in all_text)
for forbidden in ['学校名称','指导老师']:
    print('contains', forbidden, forbidden in all_text)
